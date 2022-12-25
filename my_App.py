from docx import Document 
from docx.shared import Inches
import pyttsx3
def speak(text):
        pyttsx3.speak(text)

document = Document()
#profile picture
document.add_picture('me.png')

#name phone number and email details
name = input ('what is your name?')
speak('hello ' + name + 'how are you' )
phone_number = input('what is your phone number?')
speak('your phone number is' + phone_number)

email = input('what is your email')
speak('added successfully')
document.add_paragraph(
name + ' | ' + phone_number + ' | ' + email)

#about
document.add_heading('About me')
document.add_paragraph(
    input('Tell me about youself?')
)


#work expierence
document.add_heading('work experience')
p = document.add_paragraph()

company = input('Enter company')
from_date = input('From Date')
to_date =input('To Date')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

expierence_details = input('Decribe your expierence @' + company)
p.add_run(expierence_details)

#more expierence
while True:
     has_more_expeirences = input(
        'do you have more expereinces? Yes or No '
     )
     if has_more_expeirences.lower() == 'yes':
        p = document.add_paragraph()
        

        company = input('Enter company')
        from_date = input('From Date')
        to_date =input('To Date')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        expierence_details = input('Decribe your expierence @' + company )
        p.add_run(expierence_details)
     else:
        break
#skills
document.add_heading ('Skills Acquired')
skill = input('Enter skill')
p = document.add_paragraph(skill)
p.style = 'List Bullet'   
while True:
    more_skill = input(
        'Do you have more skills? yes or no'
        )
    if more_skill.lower() == 'yes':
        skill = input('Enter skill')
        p = document.add_paragraph(skill)
    else:
        break 


